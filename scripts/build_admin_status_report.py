from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "reportes"
DOCX_PATH = OUT_DIR / "informe_avance_multi_accesorios.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "111827"
MUTED = "6B7280"
RED = "E30613"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F5F6F8"
GREEN_FILL = "EAF7EE"
AMBER_FILL = "FFF7E6"
RED_FILL = "FDEBEC"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK, size: int = 9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    header = table.rows[0].cells
    for idx, label in enumerate(headers):
        set_cell_text(header[idx], label, bold=True, color=DARK_BLUE, size=9)
        set_cell_shading(header[idx], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], str(text), size=8 if len(str(text)) > 90 else 9)
    doc.add_paragraph()
    return table


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [1.35, 2.15, 3.0])
    for idx, label in enumerate(["Estado", "Área", "Detalle"]):
        set_cell_text(table.rows[0].cells[idx], label, bold=True, color=DARK_BLUE, size=9)
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
    fill_by_status = {
        "Listo": GREEN_FILL,
        "Local": AMBER_FILL,
        "Pendiente": RED_FILL,
    }
    for status, area, detail in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], status, bold=True, color=INK, size=9)
        set_cell_shading(cells[0], fill_by_status.get(status, LIGHT_GRAY))
        set_cell_text(cells[1], area, bold=True, color=INK, size=9)
        set_cell_text(cells[2], detail, size=8 if len(detail) > 100 else 9)
    doc.add_paragraph()


def add_callout(doc, title: str, body: str, fill: str = LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [6.5])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Calibri"
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Multi Accesorios - Informe de avance y operación")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    configure_styles(doc)
    add_footer(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    mark = p.add_run("m ")
    mark.bold = True
    mark.italic = True
    mark.font.name = "Georgia"
    mark.font.size = Pt(34)
    mark.font.color.rgb = RGBColor.from_string(RED)
    brand = p.add_run("MULTI ACCESORIOS")
    brand.bold = True
    brand.font.name = "Calibri"
    brand.font.size = Pt(20)
    brand.font.color.rgb = RGBColor.from_string(INK)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("Informe de avance, operación y próximos pasos")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    sr = subtitle.add_run("Proyecto tienda online, administración, sincronización Bsale/Supabase y pagos")
    sr.font.name = "Calibri"
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor.from_string(MUTED)

    add_callout(
        doc,
        "Resumen ejecutivo",
        "La tienda ya cuenta con una experiencia marketplace más completa, catálogo conectado a Supabase, enriquecimiento de imágenes, carrito/checkout, administración de pedidos y una primera integración local de Webpay en ambiente de prueba. La regla operativa definida es clara: Bsale no se modifica desde la tienda; la app trabaja sobre Supabase y Bsale queda como fuente de sincronización.",
        LIGHT_BLUE,
    )

    add_status_table(
        doc,
        [
            ("Listo", "Diseño público", "Home, catálogo, detalle de producto, carrito y navegación fueron llevados a una estética marketplace más profesional en escritorio y móvil."),
            ("Listo", "Catálogo", "Categorías principales activas, incluyendo Vapers; productos reales desde Supabase, con imágenes principales y variantes con fallback cuando no hay match confiable."),
            ("Listo", "Administración", "Admin de productos, pedidos, descuentos y métricas adaptado a la nueva estética y preparado para revisar pedidos."),
            ("Local", "Transacciones", "Checkout local con métodos transferencia, pago al retirar, link de pago y Webpay en modo integración. Pendiente cerrar pruebas reales y push/deploy."),
            ("Local", "Webpay", "SDK Transbank agregado; endpoints de creación y confirmación implementados. El dinero real llegará al comercio configurado en Transbank cuando se usen credenciales productivas."),
            ("Pendiente", "Producción", "Credenciales productivas Webpay, datos bancarios definitivos, textos legales y prueba de flujo completo antes de habilitar pagos reales."),
        ],
    )

    doc.add_heading("1. Qué está funcionando hoy", level=1)
    add_bullets(
        doc,
        [
            "Catálogo público con filtros por categorías, ordenamiento y navegación a ofertas, nuevos, más vendidos y marcas.",
            "Diseño responsive de home, shop, carrito, detalle de producto y administración.",
            "Detalle de producto con variantes, imagen principal, galería, estado de stock y barra móvil de compra.",
            "Carrito con resumen, descuentos, datos de cliente y selección de entrega.",
            "Sincronización Bsale -> Supabase para productos, precios y stock; no hay escritura hacia Bsale desde la tienda.",
            "Flujo local de transacciones: el pedido guarda método/estado de pago, referencia y datos de entrega en Supabase.",
        ],
    )

    doc.add_heading("2. Regla crítica: Bsale y Supabase", level=1)
    add_callout(
        doc,
        "Regla acordada",
        "La tienda no debe descontar ni modificar stock directamente en Bsale. Bsale queda como sistema externo de referencia/sincronización. Supabase actúa como base operativa de la web: catálogo mostrado, pedidos, estados de pago, referencias y ajustes locales.",
        AMBER_FILL,
    )
    add_table(
        doc,
        ["Sistema", "Rol", "Qué se puede hacer"],
        [
            ["Bsale", "Fuente externa de productos/precios/stock", "Leer y sincronizar hacia Supabase. No modificar stock desde la tienda."],
            ["Supabase", "Base de la tienda", "Guardar productos, imágenes, variantes, pedidos, pagos y estados visibles en admin."],
            ["Vercel/Next.js", "Aplicación web", "Mostrar tienda, procesar checkout, iniciar pagos Webpay y registrar resultados."],
            ["Transbank/Webpay", "Procesador de pago", "Recibir pago del cliente y abonar al comercio contratado en Transbank."],
        ],
        [1.2, 2.0, 3.3],
    )

    doc.add_heading("3. Manual corto para el administrador", level=1)
    doc.add_heading("3.1 Revisión de productos", level=2)
    add_numbered(
        doc,
        [
            "Entrar al panel de administración.",
            "Abrir Productos para revisar nombre, precio, categoría, stock visible e imágenes.",
            "Si un producto no tiene imagen confiable, mantener fallback o subir una imagen manual.",
            "Evitar editar stock manualmente si se está usando Bsale como fuente de verdad, salvo corrección puntual autorizada.",
        ],
    )
    doc.add_heading("3.2 Revisión de pedidos", level=2)
    add_numbered(
        doc,
        [
            "Entrar a Pedidos.",
            "Revisar cliente, teléfono, correo, tipo de entrega y productos solicitados.",
            "Revisar el bloque de pago: método, estado y referencia.",
            "Para transferencia o link de pago, confirmar comprobante antes de preparar despacho.",
            "Para Webpay, considerar pagado solo cuando el estado indique Pagado/Webpay aprobado.",
        ],
    )
    doc.add_heading("3.3 Operación de pagos manuales", level=2)
    add_table(
        doc,
        ["Método", "Cómo se usa", "Acción del administrador"],
        [
            ["Transferencia", "Cliente confirma pedido y recibe instrucciones.", "Esperar comprobante, validar monto/referencia y preparar pedido."],
            ["Pago al retirar", "Cliente reserva y paga en tienda.", "Apartar pedido si corresponde y cobrar al entregar."],
            ["Link de pago", "Cliente solicita link manual.", "Enviar link por WhatsApp/teléfono registrado y marcar pagado tras confirmación."],
            ["Webpay", "Cliente paga en Transbank.", "Revisar estado aprobado en admin; no preparar si aparece fallido o pendiente."],
        ],
        [1.2, 2.45, 2.85],
    )

    doc.add_heading("4. Webpay: estado actual y qué falta", level=1)
    add_bullets(
        doc,
        [
            "Implementado localmente el SDK oficial transbank-sdk.",
            "Implementado endpoint para crear transacción Webpay y redirigir al cliente.",
            "Implementado endpoint de retorno para confirmar token_ws con Transbank.",
            "Si Webpay aprueba, el pedido queda como pagado y se descuenta stock local en Supabase.",
            "En ambiente de integración no se mueve dinero real; se usan credenciales y tarjetas de prueba.",
        ],
    )
    add_callout(
        doc,
        "A quién llega el dinero",
        "En producción, el dinero llega al comercio asociado al código de comercio Webpay configurado en Transbank. Debe ser el contrato de Multi Accesorios SpA y la cuenta bancaria registrada en Transbank. La app no decide la cuenta de abono.",
        GREEN_FILL,
    )
    add_table(
        doc,
        ["Falta", "Responsable sugerido", "Observación"],
        [
            ["Contrato/credenciales Webpay productivas", "Administración/contador/representante legal", "Solicitar o confirmar comercio productivo en Transbank."],
            ["Variables de entorno", "Desarrollo", "WEBPAY_ENV, WEBPAY_COMMERCE_CODE, WEBPAY_API_KEY y NEXT_PUBLIC_APP_URL."],
            ["Pruebas con tarjetas de integración", "Desarrollo + administración", "Compra aprobada, rechazada, cancelada y retorno correcto."],
            ["Políticas visibles", "Administración", "Cambios/devoluciones, privacidad, despacho, términos y contacto."],
            ["Conciliación diaria", "Administración", "Comparar pedidos pagados en Supabase con abonos/reportes Transbank."],
        ],
        [1.95, 1.65, 2.9],
    )

    doc.add_heading("5. Detallitos pendientes antes de producción", level=1)
    add_status_table(
        doc,
        [
            ("Pendiente", "Legal/operación", "Agregar textos formales de términos, privacidad, cambios/devoluciones, política de despacho y retiro en tienda."),
            ("Pendiente", "Pagos", "Completar datos bancarios definitivos para transferencia y credenciales productivas Webpay."),
            ("Pendiente", "QA móvil", "Probar en iPhone/Android reales: carrusel, checkout, teclado, scroll, botones fijos y carga de imágenes."),
            ("Pendiente", "Pedidos", "Definir estados finales: pendiente, pagado, preparado, entregado, cancelado, devuelto."),
            ("Pendiente", "Emails/WhatsApp", "Automatizar confirmación al cliente y alerta interna cuando entra pedido o pago aprobado."),
            ("Pendiente", "Deploy", "Subir cambios locales de transacciones/Webpay, validar build Vercel y probar con URL pública."),
        ],
    )

    doc.add_heading("6. Checklist de salida a producción", level=1)
    add_bullets(
        doc,
        [
            "Build de Vercel sin errores y variables de entorno cargadas.",
            "Supabase actualizado con esquema de pagos y pedidos.",
            "Webpay probado con flujo aprobado, rechazado y cancelado.",
            "Pedido Webpay aprobado limpia carrito y aparece como pagado en admin.",
            "Pedido manual aparece con referencia y método correcto.",
            "Stock local no queda negativo tras pruebas.",
            "Bsale sigue sin recibir modificaciones desde la tienda.",
            "Administrador sabe revisar pedidos y pagos antes de preparar entrega.",
        ],
    )

    doc.add_heading("7. Glosario breve", level=1)
    add_table(
        doc,
        ["Concepto", "Significado"],
        [
            ["Supabase", "Base de datos de la tienda online."],
            ["Bsale", "Sistema externo de inventario/venta usado como fuente de sincronización."],
            ["Webpay", "Pasarela de pago de Transbank para tarjetas y Redcompra."],
            ["paymentStatus", "Estado interno del pago: pendiente, pagado, fallido, link pendiente, etc."],
            ["Referencia", "Código visible para ubicar operación/pedido, útil en transferencia o conciliación."],
        ],
        [1.6, 4.9],
    )

    doc.add_paragraph()
    add_callout(
        doc,
        "Recomendación final",
        "Antes de habilitar pagos reales, hacer una prueba acompañada con el administrador: crear pedido, pagar en Webpay integración, revisar admin, revisar Supabase y confirmar que el pedido queda claro para preparar entrega.",
        LIGHT_BLUE,
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
